import os
from typing import List

import pytest
from docx import Document

from compare import compare_documents
from ingest import Ingestor


def make_docx(path: str, title: str, body_lines: List[str]):
    d = Document()
    d.add_heading(title, level=1)
    for line in body_lines:
        d.add_paragraph(line)
    d.save(path)


def test_revision_and_tolerance_mismatch(tmp_path):
    data_dir = tmp_path / "data"
    data_dir.mkdir()

    a_path = data_dir / "spec_A.docx"
    b_path = data_dir / "spec_B.docx"

    make_docx(
        str(a_path),
        "Widget Spec Rev E",
        [
            "Scope: This document controls the widget.",
            "General tolerance ±0.002",
        ],
    )

    make_docx(
        str(b_path),
        "Widget Spec Rev F",
        [
            "Scope: This document controls the widget.",
            "General tolerance ±0.0015",
        ],
    )

    # Ingest both docs explicitly to ensure caches are built
    ing = Ingestor()
    ing.ingest_file(str(a_path))
    ing.ingest_file(str(b_path))

    report = compare_documents([str(a_path), str(b_path)])

    types = [m["type"] for m in report["mismatches"]]
    assert "revision" in types
    assert "tolerance" in types

    # Find the tolerance mismatch and assert values
    tol_mismatch = next(m for m in report["mismatches"] if m["type"] == "tolerance")
    # order of A/B baseline may vary depending on compare logic; accept either
    vals = {tol_mismatch["a"]["value"], tol_mismatch["b"]["value"]}
    assert vals == {"0.002", "0.0015"} or vals == {"±0.002", "±0.0015"}

