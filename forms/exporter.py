from __future__ import annotations

import json
import os
from typing import Dict, List, Optional

from docx import Document  # type: ignore
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook  # type: ignore
from openpyxl.styles import Alignment, Border, Side, PatternFill

from forms.autofill import DraftBundle
from utils.logger import get_logger

logger = get_logger(__name__)


def export_draft(bundle: DraftBundle, out_dir: str = "exports", fmt: str = "docx") -> List[str]:
    os.makedirs(out_dir, exist_ok=True)
    outputs: List[str] = []
    if bundle.draft_type == "as9102":
        if fmt in ("docx", "all"):
            outputs.append(_export_as9102_docx(bundle, out_dir))
        if fmt in ("xlsx", "all"):
            outputs.extend(_export_as9102_xlsx(bundle, out_dir))
    elif bundle.draft_type == "8d":
        if fmt in ("docx", "all"):
            outputs.append(_export_8d_docx(bundle, out_dir))
    else:
        raise ValueError("Unknown draft type")

    # Persist bundle JSON alongside exports
    with open(os.path.join(out_dir, f"draft_{bundle.draft_id}.json"), "w", encoding="utf-8") as f:
        json.dump(bundle.__dict__, f, default=lambda o: o.__dict__, ensure_ascii=False, indent=2)
    return outputs


def _export_as9102_docx(bundle: DraftBundle, out_dir: str) -> str:
    doc = Document()
    doc.add_heading("AS9102 First Article Inspection Report — Draft", level=1)
    # Form 1 summary
    if bundle.form1:
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for label, val in [
            ("Part Number", bundle.form1.part_number),
            ("Part Name", bundle.form1.part_name),
            ("Revision", bundle.form1.part_revision),
            ("Material", bundle.form1.material),
            ("Special Processes", bundle.form1.special_processes),
        ]:
            row = t.add_row().cells
            row[0].text = label
            row[1].text = str(val or "")
    doc.add_paragraph("")

    # Evidence appendix
    doc.add_heading("Evidence & Citations", level=2)
    et = doc.add_table(rows=1, cols=4)
    hdr = et.rows[0].cells
    hdr[0].text = "Field Path"
    hdr[1].text = "Source File"
    hdr[2].text = "Page/Row"
    hdr[3].text = "Excerpt"
    for field, cits in (bundle.provenance or {}).items():
        for c in cits:
            row = et.add_row().cells
            row[0].text = field
            row[1].text = getattr(c, "filename", "")
            row[2].text = f"p{getattr(c, 'page', '')}/l{getattr(c, 'line', '')}"
            row[3].text = getattr(c, "excerpt", "")

    out = os.path.join(out_dir, f"as9102_{bundle.draft_id}.docx")
    doc.save(out)
    logger.info("Exported AS9102 DOCX to %s", out)
    return out


def _export_as9102_xlsx(bundle: DraftBundle, out_dir: str) -> List[str]:
    thin = Side(style="thin", color="DDDDDD")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)
    wrap = Alignment(wrap_text=True, vertical="top")
    fail_fill = PatternFill(start_color="FFF5F5", end_color="FFF5F5", fill_type="solid")

    outs: List[str] = []

    # Form 2 — Characteristics
    wb2 = Workbook()
    ws2 = wb2.active
    ws2.title = "Form2"
    ws2.append(["ID", "Description", "Nominal", "Tolerance", "Unit", "Citation"])
    if bundle.form2:
        for ch in bundle.form2.characteristics:
            cit = f"{ch.citation.filename} p{ch.citation.page} l{ch.citation.line}"
            ws2.append([ch.char_id, ch.description, ch.nominal, ch.tolerance, ch.unit, cit])
    for row in ws2.iter_rows(min_row=1, max_col=6, max_row=ws2.max_row):
        for cell in row:
            cell.border = border
            cell.alignment = wrap
    out2 = os.path.join(out_dir, f"as9102_form2_{bundle.draft_id}.xlsx")
    wb2.save(out2)
    outs.append(out2)

    # Form 3 — Measurements
    wb3 = Workbook()
    ws3 = wb3.active
    ws3.title = "Form3"
    ws3.append(["ID", "Description", "Nominal", "Tolerance", "Measured", "Unit", "Instrument", "Pass/Fail", "Provenance"])
    if bundle.form3:
        for m in bundle.form3.measurements:
            pf = "PASS" if m.pass_fail else ("FAIL" if m.pass_fail is not None else "")
            prov = f"{m.provenance.get('source_file','')} row {m.provenance.get('row_number','')}"
            ws3.append([m.char_id, m.description, m.nominal, m.tolerance, m.measured, m.unit, m.instrument, pf, prov])
            if m.pass_fail is False:
                for cell in ws3[ws3.max_row]:
                    cell.fill = fail_fill
    for row in ws3.iter_rows(min_row=1, max_col=9, max_row=ws3.max_row):
        for cell in row:
            cell.border = border
            cell.alignment = wrap
    out3 = os.path.join(out_dir, f"as9102_form3_{bundle.draft_id}.xlsx")
    wb3.save(out3)
    outs.append(out3)

    logger.info("Exported AS9102 XLSX to %s, %s", out2, out3)
    return outs


def _export_8d_docx(bundle: DraftBundle, out_dir: str) -> str:
    doc = Document()
    doc.add_heading("8D / CAPA Report — Draft", level=1)
    if bundle.eightd:
        fields = [
            ("D1 Team", bundle.eightd.D1_team),
            ("D2 Problem", bundle.eightd.D2_problem),
            ("D3 Containment", bundle.eightd.D3_containment),
            ("D4 Root Cause", bundle.eightd.D4_root_cause),
            ("D5 Corrective Action", bundle.eightd.D5_corrective_action),
            ("D6 Validate", bundle.eightd.D6_validate),
            ("D7 Prevent Recurrence", bundle.eightd.D7_prevent_rec),
            ("D8 Congratulate", bundle.eightd.D8_congratulate),
        ]
        t = doc.add_table(rows=0, cols=2)
        for k, v in fields:
            row = t.add_row().cells
            row[0].text = k
            row[1].text = str(v or "")

    doc.add_paragraph("")
    doc.add_heading("Evidence & Citations", level=2)
    et = doc.add_table(rows=1, cols=4)
    hdr = et.rows[0].cells
    hdr[0].text = "Field Path"
    hdr[1].text = "Source File"
    hdr[2].text = "Page/Row"
    hdr[3].text = "Excerpt"
    for field, cits in (bundle.provenance or {}).items():
        for c in cits:
            row = et.add_row().cells
            row[0].text = field
            row[1].text = getattr(c, "filename", "")
            row[2].text = f"p{getattr(c, 'page', '')}/l{getattr(c, 'line', '')}"
            row[3].text = getattr(c, "excerpt", "")

    out = os.path.join(out_dir, f"8d_{bundle.draft_id}.docx")
    doc.save(out)
    logger.info("Exported 8D DOCX to %s", out)
    return out

