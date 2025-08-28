from __future__ import annotations

import json
import os
import time
from dataclasses import asdict, dataclass, field
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple

from impact.eco_parser import EcoChange, parse_eco_text
from impact.graph import build_graph, load_graph, infer_owner
from ingest import Ingestor
from search import SearchEngine
from storage.doc_store import DocStore
from utils.hashing import sha256_hex
from utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class Citation:
    filename: str
    page: int
    line: int
    excerpt: str


@dataclass
class ImpactItem:
    asset_type: str
    id: str  # doc_id or ref
    title: str
    owner: str
    impact_score: int
    rationale: str
    suggested_action: str
    citations: List[Citation] = field(default_factory=list)
    due_date: Optional[str] = None


@dataclass
class ImpactBrief:
    id: str
    created_at: str
    eco_summary: str
    effective_date: Optional[str]
    items: List[ImpactItem]
    citations: List[Citation]
    risk_note: str
    reviewer_list: List[str]
    next_actions: str


def _find_first_citation(doc_store: DocStore, doc_id: str, needles: List[str]) -> Optional[Citation]:
    pages = doc_store.load_doc_lines(doc_id) or []
    for p_idx, lines in enumerate(pages, start=1):
        for l_idx, line in enumerate(lines, start=1):
            for n in needles:
                if n and n.lower() in (line or "").lower():
                    return Citation(filename=doc_store.docs[doc_id].filename, page=p_idx, line=l_idx, excerpt=(line or "")[:200])
    return None


def _suggest_action(asset_type: str) -> str:
    return {
        "SOP": "Review & revise",
        "CHECKLIST": "Review & revise",
        "TRAINING": "Review & revise",
        "NC_PROGRAM": "Validate/tune offsets",
        "FIXTURE": "Inspect GD&T / recertify",
        "FORM_TEMPLATE": "Update field/tolerance mapping",
    }.get(asset_type, "Review")


def _is_recent(path: Optional[str], months: int = 12) -> bool:
    try:
        if not path or not os.path.exists(path):
            return False
        mtime = os.path.getmtime(path)
        return (datetime.now() - datetime.fromtimestamp(mtime)) <= timedelta(days=30 * months)
    except Exception:
        return False


def _source_path(filename: str) -> Optional[str]:
    candidates = [os.path.join("data", "uploads", filename), os.path.join("data", filename)]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def _score_item(asset_type: str, has_direct_ref: bool, similar: bool, is_recent_mod: bool, obsolete: bool, derived: bool) -> int:
    score = 0
    if has_direct_ref:
        score += 40
    if derived:
        score += 20
    if similar:
        score += 10
    if is_recent_mod:
        score += 10
    if asset_type in {"SOP", "CHECKLIST", "TRAINING"}:
        score += 10
    if obsolete:
        score -= 10
    return max(0, min(100, score))


def _obsolete_flag(doc_store: DocStore, doc_id: str) -> bool:
    pages = doc_store.load_doc_lines(doc_id) or []
    for lines in pages:
        text = "\n".join(lines).lower()
        if any(w in text for w in ["obsolete", "superseded", "cancelled"]):
            return True
    return False


def build_brief(
    eco_text: str,
    linked_docs: Optional[List[str]] = None,
    effective_date: Optional[str] = None,
    filters: Optional[Dict[str, Any]] = None,
    max_items: int = 50,
) -> ImpactBrief:
    t0 = time.time()
    linked_docs = linked_docs or []
    filters = filters or {}
    eco = parse_eco_text(eco_text or "", effective_date=effective_date)

    ing = Ingestor()
    doc_store = ing.doc_store
    # Ensure graph exists
    graph = load_graph()
    if not graph.get("nodes"):
        graph = build_graph(doc_store)

    # Candidate docs: filter by types/owners if requested
    types_filter = set(filters.get("types", [])) if filters else set()
    owners_filter = set(filters.get("owners", [])) if filters else set()

    # Similarity search using SearchEngine
    search = SearchEngine()
    q_parts = eco.parts + [f"§{c}" for c in eco.clauses] + eco.materials + eco.deltas
    sim_docs: Dict[str, float] = {}
    if q_parts:
        q = " ".join(q_parts)
        for r in search.search(q, top_k=min(100, max_items * 3)):
            fn = r["citation"]["filename"]
            # Find doc_id by filename
            for di in doc_store.docs.values():
                if di.filename == fn:
                    sim_docs[di.doc_id] = r["score"]

    items: List[ImpactItem] = []

    # Walk doc store
    needles = eco.parts + eco.materials + eco.clauses
    for di in doc_store.docs.values():
        asset_type = infer_owner("", di.filename)  # not correct; use type from filename
        # Better: infer type quickly
        from impact.graph import infer_type

        asset_type = infer_type(di.filename)
        owner = infer_owner(asset_type, di.filename)
        if types_filter and asset_type not in types_filter:
            continue
        if owners_filter and owner not in owners_filter:
            continue

        # Direct reference check
        citation = _find_first_citation(doc_store, di.doc_id, [f"§{c}" for c in eco.clauses] + eco.parts + eco.materials)
        has_direct = citation is not None

        # Similarity (from search results)
        similar = di.doc_id in sim_docs

        # Derived heuristic: filename containing part id
        derived = any(p.replace(" ", "-") in di.filename for p in eco.parts)

        is_recent = _is_recent(_source_path(di.filename))
        obsolete = _obsolete_flag(doc_store, di.doc_id)
        score = _score_item(asset_type, has_direct, similar, is_recent, obsolete, derived)
        if score == 0:
            continue

        rationale_bits = []
        if has_direct:
            rationale_bits.append("direct clause/part/material match")
        if similar:
            rationale_bits.append("semantic similarity")
        if derived:
            rationale_bits.append("same part lineage")
        if is_recent:
            rationale_bits.append("recently modified")
        if obsolete:
            rationale_bits.append("obsolete")
        rationale = ", ".join(rationale_bits) or "related"

        items.append(
            ImpactItem(
                asset_type=asset_type,
                id=di.doc_id,
                title=di.filename,
                owner=owner,
                impact_score=score,
                rationale=rationale,
                suggested_action=_suggest_action(asset_type),
                citations=[citation] if citation else [],
            )
        )

    # Sort and cap
    items.sort(key=lambda x: x.impact_score, reverse=True)
    items = items[:max_items]

    # Risk note
    risk_bits = []
    if eco.materials:
        risk_bits.append(f"Material change: {', '.join(eco.materials)}")
    if eco.clauses:
        risk_bits.append(f"Clause change: {', '.join('§'+c for c in eco.clauses)}")
    if eco.deltas:
        risk_bits.append(f"Deltas: {', '.join(eco.deltas)}")
    risk_note = "; ".join(risk_bits) or "No high-risk changes detected"

    brief_id = sha256_hex(eco_text + (effective_date or ""))[:16]
    brief = ImpactBrief(
        id=brief_id,
        created_at=datetime.utcnow().isoformat(),
        eco_summary=(eco_text or "")[:400],
        effective_date=effective_date or eco.effective_date,
        items=items,
        citations=[c for it in items for c in it.citations],
        risk_note=risk_note,
        reviewer_list=[],
        next_actions="Review high-impact items and assign owners.",
    )

    # Persist brief
    os.makedirs("storage/impact", exist_ok=True)
    with open(os.path.join("storage/impact", f"brief_{brief.id}.json"), "w", encoding="utf-8") as f:
        json.dump(
            {
                "brief": asdict(brief),
                "eco": asdict(eco),
            },
            f,
            ensure_ascii=False,
            indent=2,
        )
    logger.info("Impact brief built id=%s items=%d duration_ms=%d", brief.id, len(items), int((time.time()-t0)*1000))
    return brief


def export_brief_docx(brief: ImpactBrief, out_dir: str = "exports") -> str:
    from docx import Document  # type: ignore

    os.makedirs(out_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Change Impact Brief", 0)
    doc.add_paragraph(f"Brief ID: {brief.id}")
    doc.add_paragraph(f"Created: {brief.created_at}")
    if brief.effective_date:
        doc.add_paragraph(f"Effective: {brief.effective_date}")
    doc.add_paragraph(brief.eco_summary)
    doc.add_paragraph(f"Risk: {brief.risk_note}")

    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Asset"
    hdr[1].text = "Owner"
    hdr[2].text = "Impact"
    hdr[3].text = "Rationale"
    hdr[4].text = "Suggested Action"
    hdr[5].text = "Citation"
    for it in brief.items:
        row = table.add_row().cells
        row[0].text = f"{it.asset_type} — {it.title}"
        row[1].text = it.owner
        row[2].text = str(it.impact_score)
        row[3].text = it.rationale
        row[4].text = it.suggested_action
        if it.citations:
            c = it.citations[0]
            row[5].text = f"{c.filename} p{c.page} l{c.line}"

    # Evidence & Citations appendix
    doc.add_page_break()
    doc.add_heading("Evidence & Citations", level=1)
    t2 = doc.add_table(rows=1, cols=4)
    h2 = t2.rows[0].cells
    h2[0].text = "Field"
    h2[1].text = "Source"
    h2[2].text = "Page/Line"
    h2[3].text = "Excerpt"
    for it in brief.items:
        for c in it.citations:
            r = t2.add_row().cells
            r[0].text = it.title
            r[1].text = c.filename
            r[2].text = f"p{c.page}/l{c.line}"
            r[3].text = c.excerpt

    out = os.path.join(out_dir, f"impact_brief_{brief.id}.docx")
    doc.save(out)
    return out

